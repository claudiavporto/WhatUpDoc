"""Generate synthetic sample documents for WhatUpDoc testing.

    python data/make_samples.py

Creates three fully fictional documents in data/raw/, one per target
domain from the project proposal:

    sample_lease_agreement.pdf     — legal contract (PDF)
    sample_medical_record.docx     — dummy medical record (DOCX)
    sample_utility_policy.txt      — public utility policy (TXT)

All names, addresses, and clinical details are invented. No real PII
or PHI is used anywhere in this project; that is the point.
"""

from __future__ import annotations

from pathlib import Path

RAW_DIR = Path(__file__).resolve().parent / "raw"
RAW_DIR.mkdir(parents=True, exist_ok=True)

LEASE_TEXT = [
    ("RESIDENTIAL LEASE AGREEMENT", True),
    ("This Residential Lease Agreement (\"Agreement\") is entered into on "
     "March 1, 2026, between Harborview Properties LLC (\"Landlord\") and "
     "Jordan T. Ellery (\"Tenant\") for the premises located at 412 Mariner's "
     "Way, Unit 2B, Norfolk, Virginia 23510 (\"Premises\").", False),
    ("1. TERM. The initial term of this lease is twelve (12) months, "
     "commencing April 1, 2026 and ending March 31, 2027. Thereafter the "
     "lease converts to month-to-month unless either party gives sixty (60) "
     "days written notice.", False),
    ("2. RENT. Tenant shall pay rent of $1,850.00 per month, due on the "
     "first (1st) day of each month. Payments received after the fifth (5th) "
     "day of the month incur a late fee equal to five percent (5%) of the "
     "monthly rent ($92.50). Returned payments incur an additional $50.00 fee.", False),
    ("3. SECURITY DEPOSIT. Tenant shall deposit $1,850.00 with Landlord as "
     "security. The deposit shall be returned within forty-five (45) days of "
     "lease termination, less lawful deductions itemized in writing.", False),
    ("4. UTILITIES. Tenant is responsible for electricity, gas, internet, "
     "and water/sewer charges. Landlord pays for trash collection and common "
     "area lighting. Tenant shall maintain utility service in Tenant's name "
     "for the duration of the tenancy.", False),
    ("5. MAINTENANCE AND REPAIRS. Tenant shall promptly notify Landlord of "
     "any condition requiring repair. Landlord shall make repairs affecting "
     "habitability within seven (7) days of written notice. Repairs "
     "necessitated by Tenant negligence shall be at Tenant's expense.", False),
    ("6. PETS. One (1) domestic cat or dog under forty (40) pounds is "
     "permitted with a non-refundable pet fee of $300.00 and additional pet "
     "rent of $35.00 per month. Exotic animals are prohibited.", False),
    ("7. EARLY TERMINATION. Tenant may terminate early upon sixty (60) days "
     "notice and payment of a termination fee equal to two (2) months' rent, "
     "except as protected by the Servicemembers Civil Relief Act.", False),
    ("8. GOVERNING LAW. This Agreement is governed by the Virginia "
     "Residential Landlord and Tenant Act, Va. Code § 55.1-1200 et seq.", False),
]

MEDICAL_RECORD = {
    "title": "PATIENT ENCOUNTER SUMMARY — FICTIONAL TRAINING RECORD",
    "sections": [
        ("Patient Information",
         "Name: Casey R. Delmar (fictional). DOB: 07/14/1968. MRN: TR-000123. "
         "Encounter date: June 3, 2026. Provider: Dr. A. Okafor, Internal Medicine."),
        ("Chief Complaint",
         "Follow-up for hypertension and type 2 diabetes management; reports "
         "intermittent dizziness on standing over the past two weeks."),
        ("Current Medications",
         "Lisinopril 20 mg once daily; Metformin 1000 mg twice daily; "
         "Atorvastatin 40 mg nightly; Aspirin 81 mg once daily."),
        ("Allergies",
         "Penicillin — hives. Sulfa drugs — rash. No known food allergies."),
        ("Vital Signs",
         "BP 128/78 seated, 112/70 standing. HR 72. BMI 29.4. "
         "Point-of-care A1c: 7.1% (down from 7.6% in January)."),
        ("Assessment and Plan",
         "1) Hypertension: orthostatic symptoms likely medication-related; "
         "reduce Lisinopril to 10 mg daily and recheck in 4 weeks. "
         "2) Type 2 diabetes: improving control; continue Metformin at current "
         "dose. 3) Preventive: schedule annual retinal exam and lipid panel. "
         "Patient counseled on hydration and slow positional changes."),
        ("Follow-Up",
         "Return visit July 1, 2026, or sooner if dizziness worsens or "
         "syncope occurs. Lab order placed for BMP prior to visit."),
    ],
}

UTILITY_POLICY = """SPRINGFIELD REGIONAL WATER AUTHORITY
CROSS-CONNECTION CONTROL AND SERVICE LATERAL POLICY (FICTIONAL SAMPLE)
Policy No. WA-2026-04 | Effective January 1, 2026

SECTION 1 — PURPOSE
This policy establishes requirements for cross-connection control and
clarifies ownership and maintenance responsibility for water service
laterals within the Authority's service area.

SECTION 2 — SERVICE LATERAL OWNERSHIP
The Authority owns and maintains the water main, the service tap, and the
lateral from the main to and including the meter and meter box. The
property owner owns and maintains the service lateral from the outlet
side of the meter to the structure, including any private valves,
pressure-reducing devices, and interior plumbing.

SECTION 3 — BACKFLOW PREVENTION
3.1 All commercial, industrial, and irrigation connections shall be
equipped with an approved backflow prevention assembly appropriate to
the degree of hazard, as determined by the Authority's cross-connection
control program.

3.2 Backflow prevention assemblies shall be tested upon installation,
after any repair or relocation, and at least once every twelve (12)
months thereafter by a certified backflow assembly tester. Test reports
shall be submitted to the Authority within ten (10) business days of
the test.

3.3 Failure to complete required testing within thirty (30) days of
written notice constitutes grounds for termination of water service
until compliance is demonstrated.

SECTION 4 — RESIDENTIAL REQUIREMENTS
Residential connections with auxiliary water sources (wells, cisterns,
rainwater systems) or in-ground irrigation shall install, at minimum, a
dual check valve assembly at the meter. The Authority may require a
higher level of protection where site conditions warrant.

SECTION 5 — ENFORCEMENT
The Director of Engineering or designee may enter premises at reasonable
times to inspect for cross-connections. Violations of this policy are
subject to service termination under Section 3.3 and civil penalties as
provided in the Authority's enabling legislation.
"""


def make_lease_pdf() -> Path:
    """Render the lease as a two-page PDF with reportlab."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    path = RAW_DIR / "sample_lease_agreement.pdf"
    styles = getSampleStyleSheet()
    story = []
    for text, is_title in LEASE_TEXT:
        style = styles["Title"] if is_title else styles["BodyText"]
        story.append(Paragraph(text, style))
        story.append(Spacer(1, 10))
    SimpleDocTemplate(str(path), pagesize=letter).build(story)
    return path


def make_medical_docx() -> Path:
    """Render the medical record as a DOCX with python-docx."""
    from docx import Document

    path = RAW_DIR / "sample_medical_record.docx"
    doc = Document()
    doc.add_heading(MEDICAL_RECORD["title"], level=1)
    for heading, body in MEDICAL_RECORD["sections"]:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(body)
    doc.save(path)
    return path


def make_policy_txt() -> Path:
    path = RAW_DIR / "sample_utility_policy.txt"
    path.write_text(UTILITY_POLICY, encoding="utf-8")
    return path


if __name__ == "__main__":
    for maker in (make_lease_pdf, make_medical_docx, make_policy_txt):
        print("Created", maker())
