# %% [markdown]
# **Document Ingestion Pipeline**

# %%
# Import Libraries
import os
from pathlib import Path

import fitz
import spacy
from docx import Document

# %% [markdown]
# ***Configuration Variables***

# %%
# Preset Variables configs for use throughout
DATA_DIRECTORY = Path("docs")

CHUNK_SIZE = 512
CHUNK_OVERLAP = 64

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}

# %%
# Initialize spaCy w/ Sentencizer for faster and more efficient retrieval (by sentence vs full doc)
nlp = spacy.blank("en")
nlp.add_pipe("sentencizer")


# %%
# Create Text Cleaning Function to remove unnecessary whitespaces but preserve punc/nums/stop words as source material
def clean_text(text):
    text = text.strip()
    text = " ".join(text.split())

    return text

# %%
# Create PDF Parsing function using PyMuPDF
def parse_pdf(pdf_path):
    # Open out PDF file from path
    doc = fitz.open(pdf_path)
    # Extract PDF Metadata
    metadata = {"filename": Path(pdf_path).name,"title": doc.metadata.get("title", ""), "author": doc.metadata.get("author", ""), "total_pages": len(doc), "file_type": "pdf"}

    print("Document Information")
    print(f"Total Pages: {len(doc)}")
    print(f"Metadata: {doc.metadata}\n")

    pages = []
    # Go through each page of PDF
    for page_num, page in enumerate(doc, start=1):

        text = clean_text(page.get_text("text"))

        # Skip blank pages
        if not text:
            continue

        pages.append({"text": text,"metadata": {"filename": metadata["filename"], "page_number": page_num}})

    doc.close()

    return pages, metadata

# %%
# Creat DOCX Parsing function using 
def parse_docx(docx_path):
    # Open out DOCX file from path
    doc = Document(docx_path)
    # Extract DOCX Metadata (Note: python-docx iterates through paragraphs instead of pages so metadata changed accordingly)
    metadata = {"filename": Path(docx_path).name,"title": doc.core_properties.title or "", "author": doc.core_properties.author or "", "total_paragraphs": len(doc.paragraphs), "file_type": "docx"}

    print("Document Information")
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    print(f"Metadata: {metadata}\n")

    paragraphs = []
    # Go through each paragraph of DOCX
    for paragraph_num, paragraph in enumerate(doc.paragraphs, start=1):

        text = clean_text(paragraph.text)

        # Skip blank pages
        if not text:
            continue

        paragraphs.append({"text": text,"metadata": {"filename": metadata["filename"], "paragraph_number": paragraph_num}, "paragraph_style": paragraph.style.name})


    return paragraphs, metadata

# %%
# Create TXT Parsing function (paragraph-based, same shape as parse_docx)
def parse_txt(txt_path):
    # Read the file, falling back to latin-1 if it's not valid UTF-8
    try:
        raw_text = Path(txt_path).read_text(encoding="utf-8")
    except UnicodeDecodeError:
        raw_text = Path(txt_path).read_text(encoding="latin-1")

    # Some source records store paragraph breaks as a literal two-character
    # "\n" sequence rather than an actual newline byte. Normalize both to
    # real newlines before splitting.
    raw_text = raw_text.replace("\\n", "\n")

    # This dataset separates paragraphs with a SINGLE newline, not a blank
    # line (\n\n) — splitting on \n\n would collapse the whole note into one
    # paragraph, which is what was happening before this fix.
    raw_paragraphs = raw_text.split("\n")

    metadata = {"filename": Path(txt_path).name, "title": "", "author": "", "total_paragraphs": len(raw_paragraphs), "file_type": "txt"}

    print("Document Information")
    print(f"Total Paragraphs: {len(raw_paragraphs)}")
    print(f"Metadata: {metadata}\n")

    paragraphs = []
    for paragraph_num, raw_paragraph in enumerate(raw_paragraphs, start=1):

        text = clean_text(raw_paragraph)

        if not text:
            continue

        paragraphs.append({"text": text, "metadata": {"filename": metadata["filename"], "paragraph_number": paragraph_num}})

    metadata["total_paragraphs"] = len(paragraphs)

    return paragraphs, metadata

# %%
# Parse Doc based on Filetype, Return Error if wrong file type used
def parse_document(file_path):
    extension = Path(file_path).suffix.lower()

    if extension == ".pdf":
        return parse_pdf(file_path)

    elif extension == ".docx":
        return parse_docx(file_path)
    
    elif extension == ".txt":
        return parse_txt(file_path)

    else:
        raise ValueError(f"Unsupported file type: {extension}")

# %% [markdown]
# *** Next we Need to Tokenize the Text with spaCy ***

# %%
def tokenize_text(text):
    doc = nlp.make_doc(text)
    return [token.text for token in doc]

def count_tokens(text):
    return len(tokenize_text(text))


# %% [markdown]
# *** Build the Chunking Section, this breaks our docs down into smaller pieces**

# %%
# Fixed Chunking Strategy
def fixed_chunking(records, chunk_size = CHUNK_SIZE, overlap = CHUNK_OVERLAP):
    # Make sure overlap isnt larger thn Chunk Size
    if overlap >= chunk_size:
        raise ValueError("Overlap too Large.")
    
    chunks = []
    chunk_index = 0

    step_size = chunk_size - overlap

    for record in records:
        text = record["text"]
        original_metadata = record["metadata"]
        tokens = tokenize_text(text)

        for start in range(0, len(tokens), step_size):
            end = start + chunk_size
            chunk_tokens = tokens[start:end]

            if not chunk_tokens:
                continue

            chunk_text = " ".join(chunk_tokens)

            chunk_metadata = original_metadata.copy()
            chunk_metadata.update({"chunk_index": chunk_index, "chunking_strategy": "fixed", "token_count": len(chunk_tokens)})
            chunks.append({"text": chunk_text, "metadata": chunk_metadata})
            
            chunk_index += 1
    return chunks

# %%
# %%
# Create Sentence Chunking Function using spaCy
def sentence_chunking(records, chunk_size=CHUNK_SIZE):

    chunks = []
    chunk_index = 0

    # Go through each parsed record
    for record in records:

        text = record["text"]
        metadata = record["metadata"]

        doc = nlp(text)

        current_chunk = []
        current_tokens = 0

        # Go through each sentence
        for sentence in doc.sents:

            sentence = sentence.text.strip()

            if not sentence:
                continue

            sentence_tokens = count_tokens(sentence)

            # Add sentence if it still fits
            if current_tokens + sentence_tokens <= chunk_size:

                current_chunk.append(sentence)
                current_tokens += sentence_tokens

            # Otherwise save current chunk and start a new one
            else:
                if current_chunk:
                    chunk_text = " ".join(current_chunk)

                    chunk_metadata = metadata.copy()
                    chunk_metadata.update({"chunk_index": chunk_index, "chunking_strategy": "sentence", "token_count": current_tokens})

                    chunks.append({"text": chunk_text, "metadata": chunk_metadata})

                    chunk_index += 1

                current_chunk = [sentence]
                current_tokens = sentence_tokens

        # Save the final chunk
        if current_chunk:

            chunk_text = " ".join(current_chunk)

            chunk_metadata = metadata.copy()
            chunk_metadata.update({"chunk_index": chunk_index,"chunking_strategy": "sentence","token_count": current_tokens})

            chunks.append({"text": chunk_text,"metadata": chunk_metadata})

            chunk_index += 1

    return chunks

# %%

# %%
# Create Paragraph Chunking Function
def paragraph_chunking(records, chunk_size=CHUNK_SIZE):

    chunks = []
    chunk_index = 0

    current_chunk = []
    current_tokens = 0
    current_chunk_metadata = None  # metadata of the FIRST record in the current chunk

    for record in records:

        paragraph = record["text"]
        metadata = record["metadata"]

        paragraph_tokens = count_tokens(paragraph)

        if current_tokens + paragraph_tokens <= chunk_size:

            if not current_chunk:
                current_chunk_metadata = metadata  # lock in metadata of the chunk's first paragraph

            current_chunk.append(paragraph)
            current_tokens += paragraph_tokens

        else:
            
            if current_chunk:
                chunk_text = "\n\n".join(current_chunk)

                chunk_metadata = current_chunk_metadata.copy()
                chunk_metadata.update({"chunk_index": chunk_index, "chunking_strategy": "paragraph", "token_count": current_tokens})

                chunks.append({"text": chunk_text, "metadata": chunk_metadata})

                chunk_index += 1

            current_chunk = [paragraph]
            current_tokens = paragraph_tokens
            current_chunk_metadata = metadata

    if current_chunk:

        chunk_text = "\n\n".join(current_chunk)

        chunk_metadata = current_chunk_metadata.copy()
        chunk_metadata.update({"chunk_index": chunk_index, "chunking_strategy": "paragraph", "token_count": current_tokens})

        chunks.append({"text": chunk_text, "metadata": chunk_metadata})

    return chunks

# %% [markdown]
# *** Chunk Strategy Selection and Total Doc Ingestion Method ***

# %%
def chunk_document(records, strategy):
    if strategy == "fixed":
        return fixed_chunking(records)

    elif strategy == "sentence":
        return sentence_chunking(records)

    elif strategy == "paragraph":
        return paragraph_chunking(records)

    else:
        raise ValueError(f"Not a valdi chunking strategy: {strategy}")

# %%
def ingest_document(file_path, strategy):
    records, metadata = parse_document(file_path)

    chunks = chunk_document(records, strategy)

    return chunks, metadata


