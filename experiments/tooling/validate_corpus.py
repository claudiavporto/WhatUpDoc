#!/usr/bin/env python3
"""
validate_corpus.py

Owner: Claudia Porto (feature/performance-testing)

Validates the WhatUpDoc document corpus (medical, legal, policy categories) before
handoff to the ingestion pipeline. Scans .pdf, .docx, and .txt files. Checks for:
  1. Corrupted / unreadable files
  2. Image-only / non-extractable PDFs (likely scanned, needs OCR)
  3. Empty or near-empty DOCX/TXT files
  4. Exact duplicate files (via content hash)
  5. Per-category and total page counts vs. proposal targets
     (DOCX/TXT files don't have a true "page count" without rendering, so their
     size is estimated at ~500 words/page and reported separately from PDF)

Requires: pymupdf, python-docx
    pip install pymupdf python-docx

Usage:
    python validate_corpus.py --data-dir data
    python validate_corpus.py --data-dir data --min-chars-per-page 20
"""

import argparse
import hashlib
import sys
from pathlib import Path

try:
    import fitz  # PyMuPDF (older/most versions expose this name)
except ImportError:
    try:
        import pymupdf as fitz  # newer PyMuPDF versions
    except ImportError:
        print("ERROR: PyMuPDF is required but could not be imported.")
        print()
        print(f"  Python interpreter in use: {sys.executable}")
        print(f"  Python version:            {sys.version}")
        print()
        print("This usually means pip installed pymupdf into a DIFFERENT Python")
        print("interpreter than the one running this script (common on Windows")
        print("with multiple Python installs). Try:")
        print()
        print(f'  "{sys.executable}" -m pip install pymupdf')
        print()
        print("Then re-run the script using that same interpreter, e.g.:")
        print(f'  "{sys.executable}" validate_corpus.py --data-dir data')
        sys.exit(1)

try:
    import docx  # python-docx
except ImportError:
    print("ERROR: python-docx is required but could not be imported.")
    print()
    print(f"  Python interpreter in use: {sys.executable}")
    print()
    print("Install it with:")
    print(f'  "{sys.executable}" -m pip install python-docx')
    sys.exit(1)


WORDS_PER_PAGE_ESTIMATE = 500  # rough heuristic for estimating DOCX/TXT page count


def sha256_of_file(path: Path, chunk_size: int = 65536) -> str:
    """Compute a SHA-256 hash of a file's contents, for duplicate detection."""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(chunk_size), b""):
            h.update(chunk)
    return h.hexdigest()


def check_pdf(path: Path, min_chars_per_page: int) -> dict:
    """
    Open a PDF and inspect it for corruption and text extractability.

    Returns a dict with keys:
        ok (bool), pages (int), pages_estimated (bool), extractable (bool), error (str or None)
    """
    result = {"ok": False, "pages": 0, "pages_estimated": False,
              "extractable": False, "error": None}
    try:
        doc = fitz.open(str(path))
        result["pages"] = doc.page_count

        if doc.page_count == 0:
            result["error"] = "Zero pages"
            doc.close()
            return result

        # Sample up to the first 5 pages (or all pages if fewer) to check
        # whether meaningful text can be extracted. A PDF that is scanned
        # images with no OCR layer will return little or no text.
        sample_size = min(5, doc.page_count)
        total_chars = 0
        for i in range(sample_size):
            page = doc.load_page(i)
            total_chars += len(page.get_text().strip())

        avg_chars_per_page = total_chars / sample_size
        result["extractable"] = avg_chars_per_page >= min_chars_per_page
        result["ok"] = True
        doc.close()
    except Exception as e:  # PyMuPDF raises various exceptions for corrupt files
        result["error"] = str(e)
    return result


def check_docx(path: Path, min_chars_per_page: int) -> dict:
    """
    Open a DOCX and inspect it for corruption and content.

    DOCX files don't carry a rendered page count, so "pages" here is an
    estimate based on word count (WORDS_PER_PAGE_ESTIMATE words/page).

    Returns a dict with keys:
        ok (bool), pages (int, estimated), pages_estimated (bool),
        extractable (bool), error (str or None)
    """
    result = {"ok": False, "pages": 0, "pages_estimated": True,
              "extractable": False, "error": None}
    try:
        d = docx.Document(str(path))

        text_parts = [p.text for p in d.paragraphs]
        # Also pull text out of tables, since contracts/policy docs often
        # put substantive content in tables (e.g., fee schedules, monitoring tables)
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)

        full_text = "\n".join(text_parts).strip()
        char_count = len(full_text)
        word_count = len(full_text.split())

        estimated_pages = max(1, round(word_count / WORDS_PER_PAGE_ESTIMATE)) if word_count > 0 else 0
        result["pages"] = estimated_pages
        result["extractable"] = char_count >= min_chars_per_page  # at least ~1 "page" worth
        result["ok"] = True

        if word_count == 0:
            result["error"] = "No extractable text (empty or content in unsupported elements)"
            result["ok"] = True  # still "opens fine", just flagged as non-extractable below

    except Exception as e:  # python-docx raises various exceptions for corrupt/invalid files
        result["error"] = str(e)
    return result


def check_txt(path: Path, min_chars_per_page: int) -> dict:
    """
    Open a plain text file and inspect it for content.

    Like DOCX, .txt has no true "page count", so it's estimated from word
    count using the same WORDS_PER_PAGE_ESTIMATE heuristic.

    Returns a dict with keys:
        ok (bool), pages (int, estimated), pages_estimated (bool),
        extractable (bool), error (str or None)
    """
    result = {"ok": False, "pages": 0, "pages_estimated": True,
              "extractable": False, "error": None}
    try:
        # Try utf-8 first, fall back to latin-1 for files with odd encodings
        # rather than failing outright.
        try:
            full_text = path.read_text(encoding="utf-8").strip()
        except UnicodeDecodeError:
            full_text = path.read_text(encoding="latin-1").strip()

        char_count = len(full_text)
        word_count = len(full_text.split())

        estimated_pages = max(1, round(word_count / WORDS_PER_PAGE_ESTIMATE)) if word_count > 0 else 0
        result["pages"] = estimated_pages
        result["extractable"] = char_count >= min_chars_per_page
        result["ok"] = True

        if word_count == 0:
            result["error"] = "Empty file"

    except Exception as e:
        result["error"] = str(e)
    return result


def check_document(path: Path, min_chars_per_page: int) -> dict:
    """Dispatch to the right checker based on file extension."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return check_pdf(path, min_chars_per_page)
    elif suffix == ".docx":
        return check_docx(path, min_chars_per_page)
    elif suffix == ".txt":
        return check_txt(path, min_chars_per_page)
    else:
        return {"ok": False, "pages": 0, "pages_estimated": False,
                "extractable": False, "error": f"Unsupported file type: {suffix}"}


def find_documents(data_dir: Path) -> dict:
    """
    Walk data_dir looking for category subfolders (e.g., medical/, legal/, policy/)
    and collect all .pdf, .docx, and .txt files under each. Falls back to a flat
    scan if no recognizable subfolders exist.
    """
    categories = {}
    subdirs = [d for d in data_dir.iterdir() if d.is_dir()]

    def collect(folder: Path):
        files = (sorted(folder.rglob("*.pdf"))
                 + sorted(folder.rglob("*.docx"))
                 + sorted(folder.rglob("*.txt")))
        return sorted(files)

    if subdirs:
        for d in subdirs:
            files = collect(d)
            if files:
                categories[d.name] = files
    else:
        files = collect(data_dir)
        if files:
            categories["all"] = files

    return categories


def main():
    parser = argparse.ArgumentParser(description="Validate the WhatUpDoc document corpus (PDF + DOCX + TXT).")
    parser.add_argument("--data-dir", type=str, default="data",
                         help="Path to the top-level data directory (default: data)")
    parser.add_argument("--min-chars-per-page", type=int, default=20,
                         help="Minimum characters (PDF: per sampled page; DOCX/TXT: total) "
                              "to count as text-extractable (default: 20)")
    args = parser.parse_args()

    data_dir = Path(args.data_dir)
    if not data_dir.exists():
        print(f"ERROR: data directory not found: {data_dir}")
        sys.exit(1)

    categories = find_documents(data_dir)
    if not categories:
        print(f"No .pdf, .docx, or .txt files found under {data_dir}. Check the path and folder structure.")
        sys.exit(1)

    hash_to_paths = {}
    total_pages = 0
    total_pages_pdf = 0
    total_pages_docx_est = 0
    total_pages_txt_est = 0
    total_files = 0
    corrupted = []
    non_extractable = []

    print("=" * 70)
    print("WHATUPDOC CORPUS VALIDATION")
    print("=" * 70)

    for category, files in categories.items():
        cat_pages = 0
        cat_ok = 0
        print(f"\n[{category}]  ({len(files)} file(s))")

        for file_path in files:
            total_files += 1
            file_hash = sha256_of_file(file_path)
            hash_to_paths.setdefault(file_hash, []).append(file_path)

            info = check_document(file_path, args.min_chars_per_page)
            suffix = file_path.suffix.lower()

            if not info["ok"]:
                corrupted.append((file_path, info["error"]))
                print(f"  [CORRUPT]     {file_path.name}  -> {info['error']}")
                continue

            cat_ok += 1
            cat_pages += info["pages"]
            total_pages += info["pages"]
            if suffix == ".docx":
                total_pages_docx_est += info["pages"]
            elif suffix == ".txt":
                total_pages_txt_est += info["pages"]
            else:
                total_pages_pdf += info["pages"]

            page_label = f"~{info['pages']} pages (est.)" if info["pages_estimated"] else f"{info['pages']} pages"

            if not info["extractable"] or info["error"]:
                non_extractable.append(file_path)
                reason = info["error"] if info["error"] else "looks scanned/image-only, may need OCR"
                print(f"  [NO TEXT?]    {file_path.name}  ({page_label}) — {reason}")
            else:
                print(f"  [OK]          {file_path.name}  ({page_label})")

        print(f"  --> {cat_ok}/{len(files)} readable, {cat_pages} pages in this category")

    # Duplicate report
    duplicates = {h: paths for h, paths in hash_to_paths.items() if len(paths) > 1}

    print("\n" + "=" * 70)
    print("SUMMARY")
    print("=" * 70)
    print(f"Total files scanned:            {total_files}")
    print(f"Total pages (PDF, actual):      {total_pages_pdf}")
    print(f"Total pages (DOCX, estimated):  {total_pages_docx_est}  "
          f"(~{WORDS_PER_PAGE_ESTIMATE} words/page)")
    print(f"Total pages (TXT, estimated):   {total_pages_txt_est}  "
          f"(~{WORDS_PER_PAGE_ESTIMATE} words/page)")
    print(f"Total pages (combined):         {total_pages}")
    print(f"Corrupted / unreadable:         {len(corrupted)}")
    print(f"Possibly scanned/no text:       {len(non_extractable)}")
    print(f"Duplicate sets found:           {len(duplicates)}")

    if duplicates:
        print("\nDUPLICATES (identical content):")
        for h, paths in duplicates.items():
            print(f"  Hash {h[:12]}...:")
            for p in paths:
                print(f"    - {p}")

    if corrupted:
        print("\nCORRUPTED FILES (fix or remove before handoff to ingestion pipeline):")
        for p, err in corrupted:
            print(f"  - {p}: {err}")

    if non_extractable:
        print("\nNON-EXTRACTABLE / LIKELY SCANNED OR EMPTY (consider OCR, fixing, or excluding):")
        for p in non_extractable:
            print(f"  - {p}")

    print("\nDone.")

    # Non-zero exit code if there are blocking issues, useful for CI / scripts
    if corrupted:
        sys.exit(2)


if __name__ == "__main__":
    main()